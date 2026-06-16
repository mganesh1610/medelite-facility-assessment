from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .schemas import AssessmentResponse


ROOT = Path(__file__).resolve().parents[2]
LOGO_PATH = ROOT / "assets" / "infinite-logo.png"

INK = RGBColor(15, 23, 42)
MUTED = RGBColor(100, 116, 139)
PRIMARY = RGBColor(30, 64, 175)
ACCENT = RGBColor(217, 119, 6)
REPORT_WIDTH = 7.0


def build_docx(assessment: AssessmentResponse) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.38)
    section.bottom_margin = Inches(0.38)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    configure_styles(document)
    add_header(document, assessment)
    add_facility_summary(document, assessment)
    add_section_label(document, "Facility Profile")
    add_profile_table(document, assessment)
    add_section_label(document, "Quality & Safety Ratings")
    add_rating_cards(document, assessment)
    add_section_label(document, "Hospitalization & ED Performance")
    add_metric_table(document, assessment)
    add_footer_summary(document, assessment)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(8.6)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def add_header(document: Document, assessment: AssessmentResponse) -> None:
    table = document.add_table(rows=1, cols=2)
    set_table_geometry(table, [2.2, 4.8])
    style_table(table, fill="FFFFFF", border="BFDBFE")

    left, right = table.rows[0].cells
    if LOGO_PATH.exists():
        paragraph = clear_cell(left)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.add_run().add_picture(str(LOGO_PATH), width=Inches(2.05))
    else:
        set_cell_text(left, "INFINITE", "Managed by MEDELITE", bold_label=True)

    paragraph = clear_cell(right)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("FACILITY ASSESSMENT SNAPSHOT")
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = INK
    meta = paragraph.add_run(f"\nState {assessment.facility.state} | CCN {assessment.facility.ccn} | CMS {assessment.facility.processing_date or 'Unknown'}")
    meta.font.size = Pt(8)
    meta.font.color.rgb = MUTED
    add_spacing(document, 3)


def add_facility_summary(document: Document, assessment: AssessmentResponse) -> None:
    table = document.add_table(rows=1, cols=2)
    set_table_geometry(table, [5.05, 1.95])
    style_table(table, fill="EFF6FF", border="BFDBFE")

    facility, score = table.rows[0].cells
    paragraph = clear_cell(facility)
    name = paragraph.add_run(assessment.facility.display_name)
    name.bold = True
    name.font.size = Pt(11)
    name.font.color.rgb = INK
    location = paragraph.add_run(f"\n{assessment.facility.location}")
    location.font.size = Pt(8.2)
    location.font.color.rgb = MUTED

    paragraph = clear_cell(score)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label = paragraph.add_run("Opportunity Score\n")
    label.bold = True
    label.font.size = Pt(7.8)
    label.font.color.rgb = MUTED
    value = paragraph.add_run(f"{assessment.opportunity.score}/100")
    value.bold = True
    value.font.size = Pt(17)
    value.font.color.rgb = ACCENT if assessment.opportunity.score >= 45 else PRIMARY
    tier = paragraph.add_run(f"\n{assessment.opportunity.tier}")
    tier.font.size = Pt(8)
    tier.font.color.rgb = INK
    add_spacing(document, 3)


def add_section_label(document: Document, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    set_table_geometry(table, [REPORT_WIDTH])
    style_table(table, fill="EFF6FF", border="BFDBFE")
    paragraph = clear_cell(table.cell(0, 0))
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(8.2)
    run.font.color.rgb = PRIMARY


def add_profile_table(document: Document, assessment: AssessmentResponse) -> None:
    rows = [
        ("EMR", value_for(assessment, "EMR"), "Certified Beds", value_for(assessment, "Census Capacity")),
        ("Current Census", value_for(assessment, "Current Census"), "Patient Type", value_for(assessment, "Type of Patient")),
        ("Previous Coverage", value_for(assessment, "Previous Coverage from Medelite"), "Medical Coverage", value_for(assessment, "Medical Coverage")),
        ("Provider Performance", value_for(assessment, "Previous Provider Performance from Medelite"), "CMS Date", assessment.facility.processing_date or "Unknown"),
    ]
    table = document.add_table(rows=0, cols=4)
    for left_label, left_value, right_label, right_value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], left_label, "", bold_label=True, label_fill=True)
        set_cell_text(cells[1], left_value, "", bold_label=False)
        set_cell_text(cells[2], right_label, "", bold_label=True, label_fill=True)
        set_cell_text(cells[3], right_value, "", bold_label=False)
    set_table_geometry(table, [1.12, 2.38, 1.12, 2.38])
    style_table(table, border="BFDBFE")
    add_spacing(document, 3)


def add_rating_cards(document: Document, assessment: AssessmentResponse) -> None:
    table = document.add_table(rows=1, cols=4)
    set_table_geometry(table, [1.75, 1.75, 1.75, 1.75])
    for cell, rating in zip(table.rows[0].cells, assessment.ratings, strict=False):
        paragraph = clear_cell(cell)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label = paragraph.add_run(rating.label + "\n")
        label.bold = True
        label.font.size = Pt(7.3)
        label.font.color.rgb = MUTED
        value = "N/A" if rating.value is None else f"{rating.value}/5"
        value_run = paragraph.add_run(value)
        value_run.bold = True
        value_run.font.size = Pt(16)
        value_run.font.color.rgb = INK
        shade_cell(cell, "FFFFFF")
    style_table(table, border="BFDBFE")
    add_spacing(document, 3)


def add_metric_table(document: Document, assessment: AssessmentResponse) -> None:
    table = document.add_table(rows=1, cols=5)
    headers = ["Metric", "Facility", "State Avg.", "National Avg.", "Unit"]
    for cell, header in zip(table.rows[0].cells, headers, strict=False):
        set_cell_text(cell, header, "", bold_label=True, label_fill=True)

    for metric in assessment.metrics:
        unit = "%" if metric.unit == "percent" else "per 1,000 resident days"
        cells = table.add_row().cells
        values = [metric.label, metric.facility_display, metric.state_display, metric.national_display, unit]
        for cell, value in zip(cells, values, strict=False):
            set_cell_text(cell, value, "", bold_label=False)
    set_table_geometry(table, [2.15, 0.95, 0.95, 1.05, 1.90])
    style_table(table, border="BFDBFE")
    add_spacing(document, 3)


def add_footer_summary(document: Document, assessment: AssessmentResponse) -> None:
    table = document.add_table(rows=1, cols=2)
    set_table_geometry(table, [4.6, 2.4])
    style_table(table, fill="F8FAFC", border="BFDBFE")

    left, right = table.rows[0].cells
    paragraph = clear_cell(left)
    title = paragraph.add_run("Action Priorities")
    title.bold = True
    title.font.size = Pt(8.3)
    title.font.color.rgb = INK
    for item in assessment.opportunity.rationale[:3]:
        run = paragraph.add_run(f"\n- {item}")
        run.font.size = Pt(7.8)
        run.font.color.rgb = INK

    paragraph = clear_cell(right)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    source = paragraph.add_run("Source\n")
    source.bold = True
    source.font.size = Pt(8.3)
    source.font.color.rgb = INK
    add_hyperlink(paragraph, assessment.facility.medicare_url, "Medicare Care Compare profile")
    meta = paragraph.add_run("\n")
    meta.font.size = Pt(2)
    quality = "No data quality warnings detected." if not assessment.data_quality else f"{len(assessment.data_quality)} data quality warning(s)"
    q_run = paragraph.add_run(quality)
    q_run.font.size = Pt(7.4)
    q_run.font.color.rgb = RGBColor(21, 128, 61) if not assessment.data_quality else ACCENT


def clear_cell(cell):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    return cell.paragraphs[0]


def set_cell_text(cell, label: str, value: str = "", *, bold_label: bool = False, label_fill: bool = False) -> None:
    paragraph = clear_cell(cell)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(label)
    run.bold = bold_label
    run.font.name = "Arial"
    run.font.size = Pt(7.9)
    run.font.color.rgb = INK
    if value:
        value_run = paragraph.add_run(f"\n{value}")
        value_run.font.name = "Arial"
        value_run.font.size = Pt(7.8)
        value_run.font.color.rgb = MUTED
    if label_fill:
        shade_cell(cell, "F8FAFC")


def set_table_geometry(table, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    total_dxa = sum(inches_to_dxa(width) for width in widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total_dxa))

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(inches_to_dxa(width)))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=False):
            cell.width = Inches(width)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(inches_to_dxa(width)))


def inches_to_dxa(value: float) -> int:
    return int(round(value * 1440))


def style_table(table, *, fill: str | None = None, border: str = "D7E3F8") -> None:
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if fill:
                shade_cell(cell, fill)
            set_cell_border(cell, border)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_spacing(document: Document, points: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(points)


def value_for(assessment: AssessmentResponse, label: str) -> str:
    for row in assessment.report_rows:
        if row.label == label:
            return row.value
    return "Not available"


def add_hyperlink(paragraph, url: str, text: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1E40AF")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
